import json
import os                           # 文件系统操作（路径、目录等）
from langchain_openai import ChatOpenAI                   # 添加模块路径用
import requests                     # HTTP 请求（调用 LLM API）
import smtplib                      # SMTP 邮件发送
from email.mime.text import MIMEText # 邮件内容格式
from email.header import Header      # 邮件主题编码
from docx import Document            # Word 文档生成
from simple_input import wea
import ast
import re                              # 正则表达式
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader  # 文档加载器（PDF/DOCX）
from langchain_text_splitters import RecursiveCharacterTextSplitter            # 文档分块工具
from langchain_core.documents import Document as LangChainDocument             # LangChain 文档对象
from langchain_huggingface import HuggingFaceEmbeddings                        # HuggingFace 嵌入模型
from langchain_community.vectorstores import FAISS                             # FAISS 向量数据库
from langchain_community.retrievers import BM25Retriever                       # BM25 检索器
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder                       # 提示词模板
from langchain_core.runnables import RunnablePassthrough                       # LCEL 透传组件
from langchain_core.output_parsers import StrOutputParser                      # 输出解析器
from langchain_core.tools import tool
from langchain_core.messages import HumanMessage, AIMessage
from langchain_classic.agents import create_tool_calling_agent, AgentExecutor
from dotenv import load_dotenv
import jieba
topk = 3
load_dotenv()
llm = ChatOpenAI(
    base_url="https://ark.cn-beijing.volces.com/api/v3",
    api_key=os.getenv("OPENAI_API_KEY"),
    model="deepseek-v3-2-251201",
    max_tokens=2048
)
LOCAL_MODEL_PATH = "./models/text2vec-base-chinese"
FAISS_INDEX_PATH = "./faiss_index"  # FAISS 索引目录
DOCUMENTS_DIR = "./knowledge_base"  # 知识库文档目录
FENCIQI_DIR = "./fenci"  # 分词目录
def zh_token(text:str):

    return jieba.lcut(text)
# 邮箱配置
EMAIL_CONFIG = {
    "smtp_server": "smtp.qq.com",
    "smtp_port": 587,
    "from_addr": "574227197@qq.com",  # 需要你填写完整邮箱
    "password": os.getenv("EMAIL_PASSWORD")
}

print("=" * 60)
print("初始化 RAG Agent 系统（真正的 FAISS 版本）")
print("=" * 60)

# 加载 Embedding 模型 (使用 LangChain)
print(f"\n[*] 加载本地模型: {LOCAL_MODEL_PATH}")
embeddings = HuggingFaceEmbeddings(model_name=LOCAL_MODEL_PATH)
print("[OK] 本地 Embedding 模型加载成功！")
@tool
def get_weather(city: str) -> str:
    """
    查询指定城市的天气情况。当用户问某城市天气、气温、下不下雨时使用。
    """
    try:
        ther = wea()
        return ther.suan(city)
    except Exception as e:
        return f"查询天气失败：{str(e)}"
@tool
def send_email(to_addr: str, subject: str, content: str) -> str:
    """
发送邮件。当用户明确要求发邮件、把内容发到某个邮箱时使用。
    """
    try:
        msg = MIMEText(content, 'plain', 'utf-8')
        msg['Subject'] = Header(subject, 'utf-8')
        msg['From'] = EMAIL_CONFIG["from_addr"]
        msg['To'] = to_addr
        
        with smtplib.SMTP(EMAIL_CONFIG["smtp_server"], EMAIL_CONFIG["smtp_port"]) as server:
            server.starttls()
            server.login(EMAIL_CONFIG["from_addr"], EMAIL_CONFIG["password"])
            server.sendmail(EMAIL_CONFIG["from_addr"], [to_addr], msg.as_string())
        
        return f"邮件已成功发送到 {to_addr}"
    except Exception as e:
        return f"邮件发送失败: {str(e)}"
@tool
def save_report_to_desktop(content: str, filename: str = "问答总结报告.docx") -> str:
    """
    把内容保存为 Word 文档到桌面。当用户要求保存、导出、整理成文档时使用。
    """
    try:
        # 获取桌面路径
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
        # 创建文档
        doc = Document()
        doc.add_heading('问答总结报告', 0)
        doc.add_paragraph(content)
        # 保存
        filepath = os.path.join(desktop, filename)
        doc.save(filepath)
        return f"报告已保存到桌面：{filepath}"
    except Exception as e:
        return f"保存失败: {str(e)}"

print("✅ 工具系统初始化完成！")

def load_fenciqi(dire):
    if not os.path.exists(dire):
        print(f"分词器目录 {dire} 不存在")
        return 

    for filename in os.listdir(dire):
        if filename.startswith('~$'):
            continue
        
        file_path = os.path.join(dire, filename)
        if os.path.isfile(file_path):
            ext = os.path.splitext(filename)[1].lower()
            try:
                if ext != '.txt':
                    continue
                jieba.load_userdict(file_path)
                print(f"   成功加载用户词典: {filename}")
            except Exception as e:
                print(f"   加载用户词典失败: {e}")
load_fenciqi(FENCIQI_DIR)
def load_documents_from_directory(directory):
    """使用 LangChain 加载目录中的 PDF 和 DOCX 文档"""
    if not os.path.exists(directory):
        print(f"⚠️ 知识库目录 {directory} 不存在")
        return None
    
    documents = []
    supported_extensions = [".pdf", ".docx", '.doc']
    
    for filename in os.listdir(directory):
        if filename.startswith('~$'):
            continue
        
        file_path = os.path.join(directory, filename)
        if os.path.isfile(file_path):
            ext = os.path.splitext(filename)[1].lower()
            if ext in supported_extensions:
                print(f"📄 正在读取文件: {filename}")
                try:
                    if ext == '.pdf':
                        loader = PyPDFLoader(file_path)
                    elif ext in ('.docx', '.doc'):
                        loader = Docx2txtLoader(file_path)
                    
                    docs = loader.load()
                    documents.extend(docs)
                    print(f"   成功读取，内容长度: {len(docs)} 页/段")
                except Exception as e:
                    print(f"   读取失败: {e}")
    
    if not documents:
        print(f"⚠️ 目录中没有找到有效的PDF或DOCX文件")
        return None
    
    print(f"\n✅ 共加载 {len(documents)} 个文档")
    return documents

print(f"\n📂 从目录加载知识库文档: {DOCUMENTS_DIR}")
documents = load_documents_from_directory(DOCUMENTS_DIR)

print("\n✂️ 正在对文档进行分块处理...")

# 预处理：将多个连续换行符统一为两个换行符
processed_docs = []
for doc in documents:
    content = doc.page_content
    # 将多个换行符统一为两个换行符
    content = re.sub(r'[\r\n]+', '\n\n', content)
    # 确保末尾有换行，保证最后一条也能被分割
    if not content.endswith('\n\n'):
        content = content.rstrip() + '\n\n'
    new_doc = LangChainDocument(page_content=content, metadata=doc.metadata)
    processed_docs.append(new_doc)

# 使用 RecursiveCharacterTextSplitter，按空行分割
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=50,  # 设小确保触发分割
    chunk_overlap=0,
    separators=["\n\n"],  # 只按空行分割
    keep_separator=False,
)

doc_chunks = text_splitter.split_documents(processed_docs)
print(f"   共生成 {len(doc_chunks)} 个文档块")

# 打印每个chunk的内容（用于调试）
for i, chunk in enumerate(doc_chunks):

    print(f"   Chunk {i+1}: {chunk.page_content[:50]}...")

# 提取文本内容用于显示
doc_texts = [doc.page_content for doc in doc_chunks]

print("\n🔍 正在创建 FAISS 向量索引...")

# 每次都从当前文档创建新索引（保持与BM25一致）
db = FAISS.from_documents(doc_chunks, embeddings)

print(f"   创建 {len(doc_texts)} 个向量")

# ========== BM25 检索器初始化 ==========
print("\n🔍 正在创建 BM25 检索器...")
bm25_retriever = BM25Retriever.from_documents(doc_chunks,preprocess_func=zh_token)
bm25_retriever.k = topk

print("✅ BM25 检索器初始化完成！")
for i, chunk in enumerate(doc_chunks):
    print(f"\n【文档块 {i+1}】")
    print(f"原文：{chunk.page_content[:50]}...")
    # 调用真实分词函数
    words = bm25_retriever.preprocess_func(chunk.page_content)
    print(f"分词结果：{words}")
# ========== 混合检索函数（RRF排序）==========
def hybrid_retrieve(query, k=topk, rrf_k=60):
    wenben = f"""
        当前的提问:{query}.
            你现在来决策当前提问问题适合用faiss向量还是bm25算法，你来分配权重，两者权重相加等于1。fai是faiss的
            权重取值[0,1]。bm是bm25算法的权重取值[0,1]。
            faiss向量检索使用场景： 用户提问无精准关键词、口语化、转述、同义改写、模糊查询，不靠字面词语匹配，依靠语义含义匹配内容；
                                 问题用词和原文用词不一致，但表达同一个意思，例如原文：入职满 1 年 5 天年假，用户问：干满一年能休几天假期；
                                 宽泛类汇总提问：公司所有福利、日常管理制度汇总、各类休假相关规定；
                                长句自然语言闲聊式提问，缺少制度专属专有名词。
                                出现以上场景：提升 FAISS 向量检索权重，降低 BM25 权重
                                
                BM25 关键词检索适用场景：提问包含专属专业名词、制度术语、固定关键词（年假、调休、考勤、出差报销、会议室预订、培训报销、厕所罚款、加班折算）；
               精准定点查询，指定具体名词 + 限定条件 例如：入职 3 年年假天数、工作日加班怎么调休、一线城市住宿报销标准；
             用户用词和文档原文关键词高度重合、字面一致，需要精准命中条款原文；
            出现以上场景：提升 BM25 权重，降低 FAISS 向量权重
           
           返回结果必须按照以下格式输出字典{{"fai":(数值)，"bm":(数值)}}
           
        """
    zidian = json.loads(llm.invoke(wenben).content)
    fai = zidian["fai"]
    bm = zidian["bm"]
    print(f"faisis:{fai}")
    print(f"bm:{bm}")
    # BM25检索
    bm25_results = bm25_retriever.invoke(query)
    
    # FAISS检索
    faiss_results = db.similarity_search(query, k=k)
    
    # 打印BM25检索结果
    print("\n" + "="*50)
    print("[BM25检索结果]")
    for idx, doc in enumerate(bm25_results, 1):
        print(f"  排名 {idx}: {doc.page_content[:100]}...")
    
    # 打印FAISS检索结果
    print("\n[FAISS检索结果]")
    for idx, doc, in enumerate(faiss_results, 1):
        print(f"  排名 {idx}: {doc.page_content[:100]}...")
    
    # 记录每个文档在各检索器中的排名
    doc_rank = {}
    
    # BM25排名（从1开始）
    for idx, doc in enumerate(bm25_results):
        if doc.page_content not in doc_rank:
            doc_rank[doc.page_content] = {'doc': doc, 'bm25_rank': idx + 1, 'faiss_rank': None}
        else:
            doc_rank[doc.page_content]['bm25_rank'] = idx + 1
    
    # FAISS排名（从1开始）
    for idx, doc, in enumerate(faiss_results):
        if doc.page_content not in doc_rank:
            doc_rank[doc.page_content] = {'doc': doc, 'bm25_rank': None, 'faiss_rank': idx + 1}
        else:
            doc_rank[doc.page_content]['faiss_rank'] = idx + 1
    
    # 计算RRF分数
    def calculate_rrf_score(item):

        score = 0.0
        if item['bm25_rank'] is not None:
            score = score + ( 1.0 / (rrf_k + item['bm25_rank']))* bm
        if item['faiss_rank'] is not None:
            score = score + (1.0 / (rrf_k + item['faiss_rank']))* fai
        return score
    
    # 按RRF分数排序
    sorted_results = sorted(doc_rank.values(), 
                          key=calculate_rrf_score, 
                          reverse=True)
    
    # 打印混合检索最终结果
    print("\n[混合检索最终结果 (RRF融合)]")
    for idx, item in enumerate(sorted_results[:k], 1):
        bm25_info = f"BM25排名: {item['bm25_rank']}" if item['bm25_rank'] else "BM25: 未命中"
        faiss_info = f"FAISS排名: {item['faiss_rank']}" if item['faiss_rank'] else "FAISS: 未命中"
        score = calculate_rrf_score(item)
        print(f"  结果 {idx} [RRF分数: {score:.4f}]")
        print(f"    {bm25_info}, {faiss_info}")
        print(f"    内容: {item['doc'].page_content[:100]}...")
    
    print("="*50 + "\n")
    
    return [item['doc'] for item in sorted_results[:k]]

retriever = db.as_retriever(search_kwargs={"k": 2})

def retrieve_docs(query):
    """检索相关文档（混合索引）"""
    retrieved_docs = hybrid_retrieve(query, k=topk)
    retrieved = []
    for doc in retrieved_docs:
        retrieved.append((doc.page_content, 1.0))
    return retrieved

print("✅ 混合检索器初始化完成！")

# 对话历史（上下文管理）
conversation_history = []

def add_to_history(role, content):
    """添加消息到对话历史"""
    conversation_history.append({"role": role, "content": content})

def format_docs(docs):
    """格式化检索到的文档"""
    return "\n\n".join([doc.page_content for doc in docs])

@tool
def search_company_policy(question: str) -> str:
    """查询公司制度知识库。当用户问年假、调休、报销、出差、考勤、福利等公司政策时使用。"""
    retrieved_docs = hybrid_retrieve(question, k=topk)
    if not retrieved_docs:
        return "未找到相关制度文档。"
    return format_docs(retrieved_docs)

agent_tools = [get_weather, send_email, save_report_to_desktop, search_company_policy]

agent_prompt = ChatPromptTemplate.from_messages([
    (
        "system",
        "你是公司智能助手。你可以：\n"
        "- 使用 search_company_policy 查询公司制度（年假、报销、出差、考勤等）\n"
        "- 使用 get_weather 查询天气\n"
        "- 使用 send_email 发送邮件\n"
        "- 使用 save_report_to_desktop 保存文档到桌面\n\n"
        "规则：\n"
        "1. 用户打招呼时直接友好回复，不要调用工具\n"
        "2. 问公司制度时先调用 search_company_policy，再基于检索结果回答\n"
        "3. 调用 send_email 时，content 必须是完整、可直接发送的正文\n"
        "4. 可以连续调用多个工具完成用户请求\n"
        "5. 回答简洁友好",
    ),
    MessagesPlaceholder("chat_history"),
    ("human", "{input}"),
    MessagesPlaceholder("agent_scratchpad"),
])

agent = create_tool_calling_agent(llm=llm, tools=agent_tools, prompt=agent_prompt)
agent_executor = AgentExecutor(
    agent=agent,
    tools=agent_tools,
    verbose=True,
    max_iterations=8,
    handle_parsing_errors=True,
)

print("✅ Agent 创建完成！")

def build_chat_history():
    """将内存对话历史转为 LangChain Message 列表"""
    messages = []
    for msg in conversation_history:
        if msg["role"] == "用户":
            messages.append(HumanMessage(content=msg["content"]))
        else:
            messages.append(AIMessage(content=msg["content"]))
    return messages

def ask_question(question, chat_history=None):
    print("\n用户问题:", question)
    print("-" * 50)
    if chat_history is None:
        chat_history = []

    result = agent_executor.invoke({
        "input": question,
        "chat_history": chat_history,
    })
    answer = result["output"]
    print(f"回答: {answer}")
    return answer

print("\n" + "=" * 60)
print("RAG Agent 知识库问答系统")
print("=" * 60)
print("输入问题开始问答，输入 '退出' 结束，输入 '清空' 清空对话历史")
print("=" * 60)


class ai:
    def yong(self, user_input):
        if user_input == "退出":
            return "再见！"

        if user_input == "清空":
            conversation_history.clear()
            return "对话历史已清空！"

        chat_history = build_chat_history()
        answer = ask_question(user_input, chat_history=chat_history)
        add_to_history("用户", user_input)
        add_to_history("助手", answer)
        return answer

class bd:
    def qie(self,pat):
        loader = Docx2txtLoader(pat)
        doce = loader.load()

        content = doce[0].page_content
        # 将多个换行符统一为两个换行符
        content = re.sub(r'[\r\n]+', '\n\n', content)
        # 确保末尾有换行，保证最后一条也能被分割
        if not content.endswith('\n\n'):
            content = content.rstrip() + '\n\n'
        new_doc = LangChainDocument(page_content=content, metadata=doce[0].metadata)

       
        doc_c = text_splitter.split_documents([new_doc])
        global db , bm25_retriever
        db.add_documents(doc_c)
        processed_docs.append(new_doc)
        CHU = text_splitter.split_documents(processed_docs)
        bm25_retriever = BM25Retriever.from_documents(CHU, preprocess_func=zh_token)
        bm25_retriever.k = topk
    def shan(self,pa):
        global db , bm25_retriever
        for doc in processed_docs:
          if doc.metadata["source"] == pa:
            processed_docs.remove(doc)
            break
        doc_chunks = text_splitter.split_documents(processed_docs)
        db = FAISS.from_documents(doc_chunks, embeddings)

        bm25_retriever = BM25Retriever.from_documents(doc_chunks, preprocess_func=zh_token)
        bm25_retriever.k = topk 
    def rufenciwendang(self,name):
        global bm25_retriever,doc_chunks
        file_path = os.path.join(FENCIQI_DIR, name)
        jieba.load_userdict(file_path)
        bm25_retriever = BM25Retriever.from_documents(doc_chunks, preprocess_func=zh_token)
        bm25_retriever.k = topk
        for i, chunk in enumerate(doc_chunks):
            print(f"\n【文档块 {i + 1}】")
            print(f"原文：{chunk.page_content[:50]}...")
            # 调用真实分词函数
            words = bm25_retriever.preprocess_func(chunk.page_content)
            print(f"分词结果：{words}")
    def rufenciciyu(self,ci):
        jieba.add_word(ci)
    def ruffencichehui(self,ci):
        jieba.del_word(ci)
    def fencishanwen(self,name):
        global bm25_retriever,doc_chunks
        file_path = os.path.join(FENCIQI_DIR, name)
        words = []
        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line:
                    words.append(line)
        for w in words:
            jieba.del_word(w)
        bm25_retriever = BM25Retriever.from_documents(doc_chunks, preprocess_func=zh_token)
        bm25_retriever.k = topk
        os.remove(file_path)
        for i, chunk in enumerate(doc_chunks):
            print(f"\n【文档块 {i + 1}】")
            print(f"原文：{chunk.page_content[:50]}...")
            # 调用真实分词函数
            words = bm25_retriever.preprocess_func(chunk.page_content)
            print(f"分词结果：{words}")